import os
import streamlit as st
import dropbox
import pdfplumber
import docx
import pandas as pd
from io import BytesIO

user = os.environ.get("DB_USERNAME")
db_name = os.environ.get("DB_NAME")

# Import LangChain untuk embeddings, vectorstore, retrieval QA, dan document model
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_community.llms import OpenAI
from langchain.docstore.document import Document

# Ambil environment variable
DROPBOX_ACCESS_TOKEN = os.environ.get("DROPBOX_ACCESS_TOKEN")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

# Pengecekan environment variable
if not DROPBOX_ACCESS_TOKEN:
    st.error("Environment variable DROPBOX_ACCESS_TOKEN belum diatur!")
    st.stop()

if not OPENAI_API_KEY:
    st.error("Environment variable OPENAI_API_KEY belum diatur!")
    st.stop()

# Inisialisasi klien Dropbox
dbx = dropbox.Dropbox(DROPBOX_ACCESS_TOKEN)

dropbox_token = os.environ.get("DROPBOX_ACCESS_TOKEN")
st.write("Debug DROPBOX_ACCESS_TOKEN:", dropbox_token)

if not dropbox_token:
    st.error("DROPBOX_ACCESS_TOKEN tidak terbaca! Periksa secrets.toml atau nama variable!")
    st.stop()

def download_file_from_dropbox(dropbox_path):
    """Mengunduh file dari Dropbox dan mengembalikan isinya sebagai byte."""
    metadata, res = dbx.files_download(dropbox_path)
    return res.content

def extract_text_from_pdf(file_bytes):
    """Ekstrak teks dari file PDF."""
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file_bytes):
    """Ekstrak teks dari file DOCX (Word)."""
    doc = docx.Document(BytesIO(file_bytes))
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_excel(file_bytes):
    """Ekstrak teks dari file Excel dengan mengonversi isi sheet pertama ke CSV."""
    with BytesIO(file_bytes) as b:
        df = pd.read_excel(b)
    text = df.to_csv(index=False)
    return text

def list_files_in_dropbox_folder(folder_path):
    """Mengambil daftar file (path) dari folder Dropbox tertentu."""
    files = []
    try:
        result = dbx.files_list_folder(folder_path)
        for entry in result.entries:
            if isinstance(entry, dropbox.files.FileMetadata):
                files.append(entry.path_lower)
    except Exception as e:
        st.error(f"Error listing files: {e}")
    return files

def load_documents_from_dropbox(folder_path):
    """
    Mengunduh file dari folder Dropbox, mengekstrak teks berdasarkan tipe file,
    dan mengembalikan list objek Document dari LangChain.
    """
    docs = []
    files = list_files_in_dropbox_folder(folder_path)
    for file_path in files:
        file_bytes = download_file_from_dropbox(file_path)
        if file_path.endswith('.pdf'):
            text = extract_text_from_pdf(file_bytes)
        elif file_path.endswith('.docx'):
            text = extract_text_from_docx(file_bytes)
        elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
            text = extract_text_from_excel(file_bytes)
        else:
            continue  # Lewati file yang tidak didukung

        if text.strip():
            docs.append(Document(page_content=text, metadata={"source": file_path}))
    return docs

def build_vectorstore(docs):
    """Membangun vector store menggunakan FAISS dari dokumen yang diberikan."""
    embedding = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    vectorstore = FAISS.from_documents(docs, embedding)
    return vectorstore

# --- Streamlit UI ---
st.title("Chatbot Dokumen Dropbox")
st.write("Chatbot ini menjawab pertanyaan berdasarkan dokumen (PDF, DOCX, Excel) yang tersimpan di Dropbox.")

# Input path folder Dropbox
dropbox_folder = st.text_input("Masukkan path folder Dropbox (contoh: /my_documents):", value="/my_documents")

# Tombol untuk memuat dokumen
if st.button("Muat Dokumen"):
    with st.spinner("Mengunduh dan memproses dokumen dari Dropbox..."):
        docs = load_documents_from_dropbox(dropbox_folder)
    if docs:
        st.success(f"{len(docs)} dokumen berhasil dimuat.")
        st.write("Contoh sumber dokumen:", [doc.metadata["source"] for doc in docs][:5])
        
        # Membangun vector store
        with st.spinner("Membangun vector store..."):
            vectorstore = build_vectorstore(docs)
        
        # Inisialisasi Retrieval QA chain
        qa = RetrievalQA.from_chain_type(
            llm=OpenAI(openai_api_key=OPENAI_API_KEY, temperature=0),
            chain_type="stuff",
            retriever=vectorstore.as_retriever()
        )
        
        st.subheader("Tanyakan pertanyaan Anda:")
        query = st.text_input("Pertanyaan:")
        if st.button("Kirim Pertanyaan") and query:
            with st.spinner("Memproses pertanyaan..."):
                answer = qa.run(query)
            st.markdown(f"**Jawaban:** {answer}")
    else:
        st.error("Tidak ada dokumen yang ditemukan di folder tersebut.")
